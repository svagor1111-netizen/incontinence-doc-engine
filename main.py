import os
import re
import uuid
from datetime import datetime
from typing import List, Optional

from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "generated")
VN_TEMPLATE_PATH = os.path.join(BASE_DIR, "MASTER_VN.docx")
ORDER_TEMPLATE_PATH = os.path.join(BASE_DIR, "MASTER_INCONTINENCE.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)

app = FastAPI(title="Incontinence Document Generator")
app.mount("/generated", StaticFiles(directory=OUTPUT_DIR), name="generated")


# -----------------------------
# SCHEMA
# -----------------------------

class Vitals(BaseModel):
    height: str
    weight: str
    blood_pressure: str
    pulse: str
    respiration: str
    temperature: str


class Diagnosis(BaseModel):
    code: str
    label: str


class EquipmentDetail(BaseModel):
    name: str
    dx: str
    medical_necessity: str


class IncontinenceOrder(BaseModel):
    underpads_chux: bool = False
    disposable_brief: bool = False
    disposable_pullup: bool = False
    absorbent_pads_liners: bool = False
    reusable_underpants: bool = False
    waterproof_mattress_cover: bool = False
    incontinence_wash: bool = False
    incontinence_cream: bool = False
    gloves: bool = False

    sex_male: bool = False
    sex_female: bool = False

    length_6_months: bool = False
    length_12_months: bool = True

    size_s: bool = False
    size_m: bool = False
    size_l: bool = False
    size_xl_xxl: bool = False


class IncontinenceRequest(BaseModel):
    mode: str = Field(default="incontinence")
    selection_mode: str = Field(default="max")
    command: Optional[str] = ""
    explicit_items: List[str] = Field(default_factory=list)

    patient_name: str
    dob: str
    age: str
    sex: str
    patient_phone: str
    patient_address: str
    insurance_id: str

    physician_name: str
    practice_name: Optional[str] = ""
    practice_address: str
    practice_phone: str
    practice_fax: str
    npi: str

    facility_name: str
    facility_address: str
    facility_phone: str

    city: str
    state: str
    zip: str

    exam_date: str
    signature_date: str

    vitals: Vitals
    diagnoses: List[Diagnosis]

    primary_diagnosis: str
    secondary_diagnoses: str

    functional_status: str
    cognitive_status: str
    ambulatory_status: str
    general_health_status: str

    equipment_list: List[str]
    equipment_details: List[EquipmentDetail]

    equipment_1: Optional[str] = ""
    equipment_2: Optional[str] = ""
    equipment_3: Optional[str] = ""
    equipment_4: Optional[str] = ""
    equipment_5: Optional[str] = ""
    equipment_6: Optional[str] = ""
    equipment_7: Optional[str] = ""
    equipment_8: Optional[str] = ""

    incontinence_order: IncontinenceOrder


# -----------------------------
# CONSTANTS
# -----------------------------

ALLOWED_ITEMS = [
    "Under Pads / Chux",
    "Disposable Brief (Diapers)",
    "Disposable Pull-Up",
    "Absorbent Pads / Liners",
    "Reusable Underpants",
    "Waterproof Mattress Cover",
    "Incontinence Wash",
    "Incontinence Cream",
    "Gloves",
]

ITEM_TO_ORDER_FIELD = {
    "Under Pads / Chux": "underpads_chux",
    "Disposable Brief (Diapers)": "disposable_brief",
    "Disposable Pull-Up": "disposable_pullup",
    "Absorbent Pads / Liners": "absorbent_pads_liners",
    "Reusable Underpants": "reusable_underpants",
    "Waterproof Mattress Cover": "waterproof_mattress_cover",
    "Incontinence Wash": "incontinence_wash",
    "Incontinence Cream": "incontinence_cream",
    "Gloves": "gloves",
}

ITEM_ALIASES = {
    "under pads / chux": "Under Pads / Chux",
    "under pads": "Under Pads / Chux",
    "underpads": "Under Pads / Chux",
    "chux": "Under Pads / Chux",
    "chux underpads": "Under Pads / Chux",
    "bed pads": "Under Pads / Chux",
    "under pads / bed pads / chux": "Under Pads / Chux",

    "disposable brief": "Disposable Brief (Diapers)",
    "disposable brief (diapers)": "Disposable Brief (Diapers)",
    "brief": "Disposable Brief (Diapers)",
    "briefs": "Disposable Brief (Diapers)",
    "diaper": "Disposable Brief (Diapers)",
    "diapers": "Disposable Brief (Diapers)",

    "disposable pull-up": "Disposable Pull-Up",
    "pull-up": "Disposable Pull-Up",
    "pull up": "Disposable Pull-Up",
    "pullups": "Disposable Pull-Up",
    "pull ups": "Disposable Pull-Up",

    "absorbent pads / liners": "Absorbent Pads / Liners",
    "absorbent pads": "Absorbent Pads / Liners",
    "liners": "Absorbent Pads / Liners",
    "pads": "Absorbent Pads / Liners",
    "shields": "Absorbent Pads / Liners",

    "reusable underpants": "Reusable Underpants",
    "reusable underwear": "Reusable Underpants",
    "underpants": "Reusable Underpants",

    "waterproof mattress cover": "Waterproof Mattress Cover",
    "mattress cover": "Waterproof Mattress Cover",
    "waterproof sheeting": "Waterproof Mattress Cover",
    "sheeting": "Waterproof Mattress Cover",

    "incontinence wash": "Incontinence Wash",
    "wash": "Incontinence Wash",

    "incontinence cream": "Incontinence Cream",
    "cream": "Incontinence Cream",
    "barrier cream": "Incontinence Cream",

    "gloves": "Gloves",
    "glove": "Gloves",
}


# -----------------------------
# HELPERS
# -----------------------------

def sanitize_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return value[:80] or "document"


def checkbox(value: bool) -> str:
    return "☑" if value else "☐"


def format_us_date(date_str: str) -> str:
    value = (date_str or "").strip()
    if not value:
        return ""
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y"):
        try:
            dt = datetime.strptime(value, fmt)
            return dt.strftime("%m/%d/%Y")
        except ValueError:
            continue
    return value


def strip_file_citations(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"", "", text)
    text = re.sub(r"\s+\.", ".", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def clean_text(text: str) -> str:
    return strip_file_citations(text or "")


def canonicalize_item_name(value: str) -> str:
    cleaned = clean_text(value).strip()
    if not cleaned:
        return ""
    if cleaned in ALLOWED_ITEMS:
        return cleaned

    lowered = cleaned.lower()
    if lowered in ITEM_ALIASES:
        return ITEM_ALIASES[lowered]

    return ""


def order_field_to_item(field_name: str) -> Optional[str]:
    for item_name, mapped_field in ITEM_TO_ORDER_FIELD.items():
        if mapped_field == field_name:
            return item_name
    return None


def determine_selection_mode(payload: IncontinenceRequest) -> str:
    candidates = [
        clean_text(payload.selection_mode).lower(),
        clean_text(payload.command).lower(),
        clean_text(payload.mode).lower(),
    ]

    for value in candidates:
        if value in {"lvn", "list_only", "list-only", "list"}:
            return "list_only"
        if value in {"vnm", "max", "incontinence"}:
            return "max"

    return "max"


def collect_explicit_list_items(payload: IncontinenceRequest) -> List[str]:
    candidates: List[str] = []

    # Explicit list field
    candidates.extend(payload.explicit_items)

    # Allow prompt/tool to pass list through normal equipment fields too
    candidates.extend(payload.equipment_list)
    candidates.extend([d.name for d in payload.equipment_details])

    for value in [
        payload.equipment_1,
        payload.equipment_2,
        payload.equipment_3,
        payload.equipment_4,
        payload.equipment_5,
        payload.equipment_6,
        payload.equipment_7,
        payload.equipment_8,
    ]:
        if value:
            candidates.append(value)

    # Also respect any explicitly checked items from incoming order
    order = payload.incontinence_order
    for field_name in ITEM_TO_ORDER_FIELD.values():
        if getattr(order, field_name, False):
            item_name = order_field_to_item(field_name)
            if item_name:
                candidates.append(item_name)

    normalized: List[str] = []
    seen = set()
    for value in candidates:
        item = canonicalize_item_name(value)
        if item and item not in seen:
            normalized.append(item)
            seen.add(item)

    return normalized


def normalize_equipment(payload: IncontinenceRequest) -> List[str]:
    selection_mode = determine_selection_mode(payload)

    # LIST ONLY MODE: use only explicit items, do not auto-add
    if selection_mode == "list_only":
        return collect_explicit_list_items(payload)

    # MAX MODE: preserve current logic
    candidates: List[str] = []

    candidates.extend([x for x in payload.equipment_list if x in ALLOWED_ITEMS])
    candidates.extend([d.name for d in payload.equipment_details if d.name in ALLOWED_ITEMS])

    for value in [
        payload.equipment_1,
        payload.equipment_2,
        payload.equipment_3,
        payload.equipment_4,
        payload.equipment_5,
        payload.equipment_6,
        payload.equipment_7,
        payload.equipment_8,
    ]:
        if value in ALLOWED_ITEMS:
            candidates.append(value)

    order = payload.incontinence_order
    for field_name in ITEM_TO_ORDER_FIELD.values():
        if getattr(order, field_name, False):
            item_name = order_field_to_item(field_name)
            if item_name:
                candidates.append(item_name)

    if "Under Pads / Chux" not in candidates:
        candidates.insert(0, "Under Pads / Chux")

    seen = set()
    normalized = []
    for item in candidates:
        if item not in seen:
            normalized.append(item)
            seen.add(item)

    return normalized


def synced_order(payload: IncontinenceRequest, items: List[str]) -> IncontinenceOrder:
    order = payload.incontinence_order.model_copy(deep=True)

    for item, field_name in ITEM_TO_ORDER_FIELD.items():
        setattr(order, field_name, item in items)

    sex = clean_text(payload.sex).lower()
    order.sex_male = sex == "male"
    order.sex_female = sex == "female"

    order.length_6_months = False
    order.length_12_months = True
    return order


def build_equipment_details(payload: IncontinenceRequest, items: List[str]) -> List[EquipmentDetail]:
    by_name = {}
    for d in payload.equipment_details:
        canonical = canonicalize_item_name(d.name)
        if canonical in ALLOWED_ITEMS:
            by_name[canonical] = d

    result = []
    for item in items:
        if item in by_name:
            result.append(
                EquipmentDetail(
                    name=item,
                    dx=clean_text(by_name[item].dx),
                    medical_necessity=clean_text(by_name[item].medical_necessity),
                )
            )
        else:
            result.append(
                EquipmentDetail(
                    name=item,
                    dx=clean_text(payload.primary_diagnosis),
                    medical_necessity="Medically necessary for management of incontinence-related hygiene needs, MRADL limitation, skin protection, and caregiver-assisted care.",
                )
            )
    return result


def diagnosis_pairs(payload: IncontinenceRequest) -> List[tuple[str, str]]:
    pairs: List[tuple[str, str]] = []
    for dx in payload.diagnoses:
        code = clean_text(dx.code)
        label = clean_text(dx.label)
        if code or label:
            pairs.append((code, label))
    return pairs


def vn_primary_diagnosis_string(payload: IncontinenceRequest) -> str:
    pairs = diagnosis_pairs(payload)
    if pairs:
        code, label = pairs[0]
        return f"{code} {label}".strip()
    return clean_text(payload.primary_diagnosis)


def vn_secondary_diagnoses_string(payload: IncontinenceRequest) -> str:
    pairs = diagnosis_pairs(payload)
    if len(pairs) > 1:
        return "; ".join(f"{code} {label}".strip() for code, label in pairs[1:])
    return clean_text(payload.secondary_diagnoses)


def order_primary_diagnosis_string(payload: IncontinenceRequest) -> str:
    pairs = diagnosis_pairs(payload)
    if pairs:
        code, label = pairs[0]
        return f"{code} {label}".strip()
    return clean_text(payload.primary_diagnosis)


def order_secondary_diagnoses_string(payload: IncontinenceRequest) -> str:
    pairs = diagnosis_pairs(payload)
    if len(pairs) > 1:
        codes = [code for code, _ in pairs[1:] if code]
        return ", ".join(codes)
    return clean_text(payload.secondary_diagnoses)


def text_has_incontinence(text: str) -> bool:
    value = clean_text(text).lower()
    phrases = [
        "urinary incontinence",
        "fecal incontinence",
        "bowel/bladder incontinence",
        "bladder incontinence",
        "bowel incontinence",
        "incontinence",
    ]
    return any(p in value for p in phrases)


def has_documented_incontinence(payload: IncontinenceRequest) -> bool:
    for dx in payload.diagnoses:
        if text_has_incontinence(dx.label) or text_has_incontinence(dx.code):
            return True

    if text_has_incontinence(payload.primary_diagnosis):
        return True
    if text_has_incontinence(payload.secondary_diagnoses):
        return True
    if text_has_incontinence(payload.functional_status):
        return True
    if text_has_incontinence(payload.general_health_status):
        return True

    for detail in payload.equipment_details:
        if text_has_incontinence(detail.dx) or text_has_incontinence(detail.medical_necessity):
            return True

    return False


def vn_practice_address_value(payload: IncontinenceRequest) -> str:
    return clean_text(payload.practice_address)


def order_address_value(payload: IncontinenceRequest) -> str:
    practice_address = clean_text(payload.practice_address)
    if practice_address:
        return practice_address
    return clean_text(payload.patient_address) or clean_text(payload.facility_address)


def source_dx_codes_for_dme(payload: IncontinenceRequest) -> str:
    pairs = diagnosis_pairs(payload)
    if not pairs:
        return clean_text(payload.primary_diagnosis)

    codes = []
    for code, _ in pairs:
        if code and code not in codes:
            codes.append(code)

    key_support = []
    for preferred in ["F03.90", "R60.9", "R06.02"]:
        if preferred in codes and preferred not in key_support:
            key_support.append(preferred)

    selected = [codes[0]]
    for code in key_support:
        if code not in selected:
            selected.append(code)

    return ", ".join(selected)


def dme_diagnosis_line(payload: IncontinenceRequest) -> str:
    source_codes = source_dx_codes_for_dme(payload)
    if source_codes:
        return f"Source Dx: {source_codes}"
    return "Source Dx: documented diagnoses"


def dme_medical_necessity(item_name: str) -> str:
    if item_name == "Under Pads / Chux":
        return (
            "Patient has documented bowel and/or bladder incontinence requiring absorbent underpads "
            "to maintain hygiene, protect bedding and seating surfaces, and reduce risk of skin breakdown."
        )
    if item_name == "Disposable Brief (Diapers)":
        return (
            "Patient is unable to toilet independently due to cognitive impairment, weakness, and mobility limitation, "
            "requiring full absorbent briefs for containment and hygiene management."
        )
    if item_name == "Disposable Pull-Up":
        return (
            "Patient participates partially in toileting but remains incontinent and requires absorbent pull-up garments "
            "for day-to-day continence management and hygiene protection."
        )
    if item_name == "Absorbent Pads / Liners":
        return (
            "Patient has leakage episodes requiring absorbent pads or liners for hygiene maintenance and clothing protection."
        )
    if item_name == "Reusable Underpants":
        return (
            "Patient requires reusable protective undergarments for continence support and hygiene management."
        )
    if item_name == "Waterproof Mattress Cover":
        return (
            "Patient requires mattress protection due to ongoing incontinence, moisture exposure risk, and dependence "
            "in hygiene care, in order to protect bedding and maintain sanitary conditions."
        )
    if item_name == "Incontinence Wash":
        return (
            "Patient requires caregiver-assisted cleaning after incontinence episodes to maintain hygiene, "
            "protect skin integrity, and reduce infection risk."
        )
    if item_name == "Incontinence Cream":
        return (
            "Patient is at risk for skin breakdown due to chronic moisture exposure from incontinence and requires "
            "barrier cream for skin protection."
        )
    if item_name == "Gloves":
        return (
            "Caregiver provides toileting and hygiene assistance and requires gloves for infection control and safe handling."
        )
    return "Medically necessary for management of incontinence-related hygiene needs, MRADL limitation, skin protection, and caregiver-assisted care."


def normalize_details_for_vn(details: List[EquipmentDetail], payload: IncontinenceRequest) -> List[EquipmentDetail]:
    normalized: List[EquipmentDetail] = []
    dx_line = dme_diagnosis_line(payload)

    for detail in details:
        normalized.append(
            EquipmentDetail(
                name=detail.name,
                dx=dx_line,
                medical_necessity=dme_medical_necessity(detail.name),
            )
        )
    return normalized


def equipment_block_lines(details: List[EquipmentDetail]) -> List[str]:
    lines = []
    for i, detail in enumerate(details, start=1):
        block = (
            f"{i}. {clean_text(detail.name)}\n"
            f"Diagnosis: {clean_text(detail.dx)}\n"
            f"Medical Necessity: {clean_text(detail.medical_necessity)}"
        )
        lines.append(block)

    while len(lines) < 8:
        lines.append("")

    return lines[:8]


def replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = paragraph.text
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    if new_text != full_text:
        paragraph.text = new_text


def replace_text_in_doc(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def replace_line_containing(doc: Document, needle: str, new_text: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if needle in paragraph.text:
                        paragraph.text = new_text


# -----------------------------
# TEMPLATE RENDER
# -----------------------------

def fill_vn_template(payload: IncontinenceRequest, details: List[EquipmentDetail], file_id: str) -> str:
    if not os.path.exists(VN_TEMPLATE_PATH):
        raise FileNotFoundError(f"VN template not found: {VN_TEMPLATE_PATH}")

    doc = Document(VN_TEMPLATE_PATH)
    vn_details = normalize_details_for_vn(details, payload)
    equipment_lines = equipment_block_lines(vn_details)

    replacements = {
        "{{physician_name}}": clean_text(payload.physician_name),
        "{{practice_address}}": vn_practice_address_value(payload),
        "{{practice_phone}}": clean_text(payload.practice_phone),
        "{{practice_fax}}": clean_text(payload.practice_fax),
        "{{exam_date}}": format_us_date(payload.exam_date),
        "{{patient_name}}": clean_text(payload.patient_name),
        "{{dob}}": clean_text(payload.dob),
        "{{age}}": clean_text(payload.age),
        "{{sex}}": clean_text(payload.sex),
        "{{facility_name}}": clean_text(payload.facility_name),
        "{{facility_address}}": clean_text(payload.facility_address),
        "{{facility_phone}}": clean_text(payload.facility_phone),
        "{{height}}": clean_text(payload.vitals.height),
        "{{weight}}": clean_text(payload.vitals.weight),
        "{{blood_pressure}}": clean_text(payload.vitals.blood_pressure),
        "{{pulse}}": clean_text(payload.vitals.pulse),
        "{{respiration}}": clean_text(payload.vitals.respiration),
        "{{temperature}}": clean_text(payload.vitals.temperature),
        "{{primary_diagnosis}}": vn_primary_diagnosis_string(payload),
        "{{secondary_diagnoses}}": vn_secondary_diagnoses_string(payload),
        "{{functional_status}}": clean_text(payload.functional_status),
        "{{cognitive_status}}": clean_text(payload.cognitive_status),
        "{{ambulatory_status}}": clean_text(payload.ambulatory_status),
        "{{general_health_status}}": clean_text(payload.general_health_status),
        "{{equipment_1}}": equipment_lines[0],
        "{{equipment_2}}": equipment_lines[1],
        "{{equipment_3}}": equipment_lines[2],
        "{{equipment_4}}": equipment_lines[3],
        "{{equipment_5}}": equipment_lines[4],
        "{{equipment_6}}": equipment_lines[5],
        "{{equipment_7}}": equipment_lines[6],
        "{{equipment_8}}": equipment_lines[7],
        "{{signature_date}}": format_us_date(payload.signature_date),
    }

    replace_text_in_doc(doc, replacements)

    filename = f"{sanitize_filename(payload.patient_name)}_{file_id}_VN.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


def fill_order_template(payload: IncontinenceRequest, order: IncontinenceOrder, file_id: str) -> str:
    if not os.path.exists(ORDER_TEMPLATE_PATH):
        raise FileNotFoundError(f"Order template not found: {ORDER_TEMPLATE_PATH}")

    doc = Document(ORDER_TEMPLATE_PATH)

    replacements = {
        "{{patient_name}}": clean_text(payload.patient_name),
        "{{dob}}": clean_text(payload.dob),
        "{{insurance_id}}": clean_text(payload.insurance_id),
        "{{height}}": clean_text(payload.vitals.height),
        "{{weight}}": clean_text(payload.vitals.weight),
        "{{primary_diagnosis}}": order_primary_diagnosis_string(payload),
        "{{secondary_diagnoses}}": order_secondary_diagnoses_string(payload),
        "{{physician_name}}": clean_text(payload.physician_name),
        "{{practice_address}}": order_address_value(payload),
        "{{city}}": clean_text(payload.city),
        "{{state}}": clean_text(payload.state),
        "{{zip}}": clean_text(payload.zip),
        "{{practice_phone}}": clean_text(payload.practice_phone),
        "{{practice_fax}}": clean_text(payload.practice_fax),
        "{{npi}}": clean_text(payload.npi),
        "{{signature_date}}": format_us_date(payload.signature_date),
    }

    replace_text_in_doc(doc, replacements)

    replace_line_containing(
        doc,
        "Male ☐",
        f"Male {checkbox(order.sex_male)}   Female {checkbox(order.sex_female)}",
    )

    replace_line_containing(
        doc,
        "Length of Need:",
        f"Length of Need: 6 Months {checkbox(order.length_6_months)}   12 Months {checkbox(order.length_12_months)}",
    )

    replace_line_containing(
        doc,
        "Disposable Brief (Diapers)",
        f"{checkbox(order.disposable_brief)} Disposable Brief (Diapers)      {checkbox(order.size_s)} S   {checkbox(order.size_m)} M   {checkbox(order.size_l)} L   {checkbox(order.size_xl_xxl)} XL-XXL",
    )

    replace_line_containing(
        doc,
        "Disposable Pull-Up",
        f"{checkbox(order.disposable_pullup)} Disposable Pull-Up              {checkbox(order.size_s)} S   {checkbox(order.size_m)} M   {checkbox(order.size_l)} L   {checkbox(order.size_xl_xxl)} XL-XXL",
    )

    replace_line_containing(
        doc,
        "Under Pads / Chux",
        f"{checkbox(order.underpads_chux)} Under Pads / Chux               (Up to 120/month)",
    )

    replace_line_containing(
        doc,
        "Absorbent Pads / Liners",
        f"{checkbox(order.absorbent_pads_liners)} Absorbent Pads / Liners         (Up to 300/month)",
    )

    replace_line_containing(
        doc,
        "Reusable Underpants",
        f"{checkbox(order.reusable_underpants)} Reusable Underpants             {checkbox(order.size_s)} S   {checkbox(order.size_m)} M   {checkbox(order.size_l)} L   {checkbox(order.size_xl_xxl)} XL-XXL",
    )

    replace_line_containing(
        doc,
        "Waterproof Mattress Cover",
        f"{checkbox(order.waterproof_mattress_cover)} Waterproof Mattress Cover       (2/year)",
    )

    replace_line_containing(
        doc,
        "Incontinence Wash",
        f"{checkbox(order.incontinence_wash)} Incontinence Wash",
    )

    replace_line_containing(
        doc,
        "Incontinence Cream",
        f"{checkbox(order.incontinence_cream)} Incontinence Cream",
    )

    replace_line_containing(
        doc,
        "Gloves",
        f"{checkbox(order.gloves)} Gloves",
    )

    filename = f"{sanitize_filename(payload.patient_name)}_{file_id}_ORDER.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


# -----------------------------
# ROUTES
# -----------------------------

@app.get("/")
def root():
    return {"status": "incontinence backend running"}


@app.get("/health")
def health():
    return {"ok": True, "service": "incontinence-doc-engine"}


@app.post("/create_dme_documents")
def create_dme_documents(payload: IncontinenceRequest):
    current_mode = determine_selection_mode(payload)

    if clean_text(payload.mode).lower() not in {"", "incontinence", "vnm", "lvn", "max", "list_only"}:
        return {"error": "Only incontinence mode is supported."}

    if not os.path.exists(VN_TEMPLATE_PATH):
        return {"error": "MASTER_VN.docx not found in backend root."}

    if not os.path.exists(ORDER_TEMPLATE_PATH):
        return {"error": "MASTER_INCONTINENCE.docx not found in backend root."}

    items = normalize_equipment(payload)

    if current_mode == "list_only" and not items:
        return {"error": "No item list provided for LIST ONLY mode."}

    details = build_equipment_details(payload, items)
    order = synced_order(payload, items)

    file_id = datetime.utcnow().strftime("%Y%m%d%H%M%S") + "_" + uuid.uuid4().hex[:8]

    vn_path = fill_vn_template(payload, details, file_id)
    order_path = fill_order_template(payload, order, file_id)

    base_url = os.getenv("PUBLIC_BASE_URL", "https://incontinence-doc-engine.onrender.com")
    vn_url = f"{base_url}/generated/{os.path.basename(vn_path)}"
    order_url = f"{base_url}/generated/{os.path.basename(order_path)}"

    return {
        "status": "success",
        "selection_mode": current_mode,
        "vn_docx": vn_url,
        "order_docx": order_url,
        "equipment_list": items,
    }
